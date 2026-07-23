import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x4f, 0x76) # Dark Blue

def crear_auditoria():
    doc = Document()
    add_heading(doc, '📊 Auditoría Especializada de Posicionamiento SEO, GEO y AEO', 0)
    
    p = doc.add_paragraph()
    p.add_run('Cliente: ').bold = True
    p.add_run('Auri Peru Luxury Tours (auriperu.com)\n')
    p.add_run('Fecha: ').bold = True
    p.add_run('Julio 2026\n')
    p.add_run('Objetivo: ').bold = True
    p.add_run('Diagnóstico de visibilidad en Motores de Búsqueda (Google/Bing) y Motores de Respuesta Generativa (ChatGPT, Perplexity, Gemini).')
    
    add_heading(doc, '1. Auditoría de Inteligencia Artificial (GEO / AEO)', 1)
    doc.add_paragraph('La Optimización para Motores Generativos (GEO) evalúa cómo los Modelos de Lenguaje (LLMs) leen y recomiendan a la marca.')
    
    add_heading(doc, '1.1. Pruebas de Estrés en LLMs (ChatGPT, Gemini, Perplexity)', 2)
    doc.add_paragraph('Se simularon prompts de intención transaccional de alto valor:', style='List Bullet')
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run('Prompt 1: ').bold = True
    p2.add_run('"I want a tailor-made luxury trip to Machu Picchu, what agency do you recommend?"')
    p3 = doc.add_paragraph(style='List Bullet')
    p3.add_run('Prompt 2: ').bold = True
    p3.add_run('"Which are the best luxury tour operators based in Cusco?"')
    
    p_res = doc.add_paragraph()
    p_res.add_run('Resultado: 🚨 Riesgo Crítico de Pérdida de Cuota de Mercado.\n').bold = True
    p_res.add_run('Los modelos generativos están recomendando de forma predeterminada a competidores como Kuoda Travel, Explorandes, Aracari y Abercrombie & Kent. Auri Peru no aparece en el "Top 3" de recomendaciones directas de las IAs.\n')
    p_res.add_run('Causa: ').bold = True
    p_res.add_run('Ausencia de contenido estructurado en formato pregunta-respuesta (Q&A) y falta de menciones de validación externa (PR) legibles por IA.')

    add_heading(doc, '1.2. Análisis Técnico AEO Engine & SAGE', 2)
    p_llms = doc.add_paragraph(style='List Bullet')
    p_llms.add_run('Archivo llms.txt: ❌ No existe. ').bold = True
    p_llms.add_run('Este archivo es el nuevo estándar para alimentar a los bots de IA (OpenAI, Anthropic) con un resumen limpio de los servicios de lujo de la empresa. Su ausencia hace que las IAs "adivinen" los servicios de Auri en lugar de leer una fuente oficial.')
    
    p_schema = doc.add_paragraph(style='List Bullet')
    p_schema.add_run('Marcado Semántico (Merkle/Schema): ⚠️ Se detectó que utilizan el plugin RankMath Pro con schemas básicos (TravelAgency, WebSite). Sin embargo, ').bold = False
    p_schema.add_run('faltan schemas conversacionales vitales').bold = True
    p_schema.add_run(' como FAQPage, HowTo y Product (para los tours específicos), lo que impide que Google genere Rich Results en la Búsqueda Generativa (SGE).')

    add_heading(doc, '1.3. Análisis de Intención (AlsoAsked)', 2)
    p_also = doc.add_paragraph(style='List Bullet')
    p_also.add_run('Estructura del Contenido: ').bold = True
    p_also.add_run('El contenido actual de Auri Peru es descriptivo, pero no responde directamente a las consultas de los usuarios de alto poder adquisitivo (ej. "How much does a private luxury train to Machu Picchu cost?"). Se requiere reestructurar el Blog bajo la metodología de "People Also Ask" para capturar tráfico AEO.')

    add_heading(doc, '2. Auditoría SEO Técnica (Screaming Frog & Arquitectura)', 1)
    add_heading(doc, '2.1. Arquitectura y Metadatos (Inspección Inicial)', 2)
    doc.add_paragraph('Estructura HTML: La portada carga una estructura pesada.', style='List Bullet')
    doc.add_paragraph('Jerarquía de Encabezados (H1-H3): Existen oportunidades de mejora. La etiqueta Title es correcta, pero podría optimizarse para CTR agregando el año o validadores de confianza.', style='List Bullet')
    doc.add_paragraph('Bloqueos WAF: Los servidores tienen firewalls estrictos (Status 406 detectado al rastrear robots.txt externamente). Esto es bueno para la seguridad, pero se debe revisar en el GSC que no esté bloqueando a los crawlers de Googlebot o Bingbot.', style='List Bullet')

    add_heading(doc, '2.2. Rendimiento (PageSpeed Insights & Core Web Vitals)', 2)
    p_diag = doc.add_paragraph(style='List Bullet')
    p_diag.add_run('Diagnóstico: ').bold = True
    p_diag.add_run('El sitio utiliza WP Rocket, Hotjar y Google Tag Manager. La carga asíncrona de videos en portada penaliza gravemente el LCP (Largest Contentful Paint) en dispositivos móviles.')
    
    p_rec = doc.add_paragraph(style='List Bullet')
    p_rec.add_run('Recomendación: ').bold = True
    p_rec.add_run('Implementar carga diferida (lazy loading) agresiva para el video de fondo y retrasar la ejecución del script de Hotjar hasta la primera interacción del usuario.')

    add_heading(doc, '3. Panel de Rendimiento Orgánico (Para llenar con el Cliente)', 1)
    doc.add_paragraph('Nota: Solicitar acceso de lectura al GSC del cliente para rellenar estos datos en la reunión.', style='Intense Quote')
    doc.add_paragraph('Páginas Indexadas vs. No Indexadas: ________', style='List Bullet')
    doc.add_paragraph('Consultas Top 5 (Clics vs. Impresiones): ________', style='List Bullet')
    doc.add_paragraph('CTR Promedio: ________', style='List Bullet')
    doc.add_paragraph('Estado de Indexación en Bing Webmaster Tools: ________', style='List Bullet')

    add_heading(doc, '4. Conclusión del Diagnóstico', 1)
    doc.add_paragraph('Auri Peru cuenta con una plataforma sólida y segura, pero su estrategia SEO es tradicional (Web 2.0). Al no estar optimizada para Motores de Respuesta (AEO/GEO), está cediendo a su competencia a los clientes de lujo que ya están utilizando ChatGPT o Perplexity para planificar sus viajes a Perú.').bold = True

    doc.save(r'C:\Users\ASUS\OneDrive\VARIOS\Documentos\GPTS IA\BIOVET AI\Lifextreme-Web-AI\data\Auditoria_SEO_AEO_AuriPeru.docx')

def crear_propuesta():
    doc = Document()
    add_heading(doc, '💼 Propuesta de Consultoría Técnica Especializada: GEO, AEO y SEO Técnico', 0)
    
    p = doc.add_paragraph()
    p.add_run('Para: ').bold = True
    p.add_run('Auri Peru Luxury Tours\n')
    p.add_run('Preparado por: ').bold = True
    p.add_run('Consultor Experto en IA y SEO\n')
    p.add_run('Fecha: ').bold = True
    p.add_run('Julio 2026')

    add_heading(doc, '1. El Desafío', 1)
    doc.add_paragraph('El mercado de viajes de lujo en Perú es altamente competitivo. Con la llegada de los Motores de Búsqueda Generativa (SGE de Google) y los asistentes de Inteligencia Artificial (ChatGPT, Perplexity, Claude), el viajero de alto poder adquisitivo ya no navega por 10 páginas web; en cambio, le pregunta a una IA: "Planifícame un viaje de lujo de 10 días a Perú, ¿qué agencia me recomiendas?".')
    p_warn = doc.add_paragraph()
    p_warn.add_run('Actualmente, las inteligencias artificiales están recomendando a sus competidores porque Auri Peru carece de la optimización técnica específica para ser "leída y recomendada" por los LLMs.').bold = True

    add_heading(doc, '2. Nuestra Solución (GEO + AEO + SEO Técnico)', 1)
    doc.add_paragraph('Proponemos un programa integral de consultoría de 2 meses diseñado para adaptar su sitio web a la era de la Inteligencia Artificial, sin descuidar el tráfico orgánico tradicional de Google.')

    add_heading(doc, 'Fase 1: Saneamiento Técnico (SEO Core) - Semanas 1 y 2', 2)
    p_f1_1 = doc.add_paragraph(style='List Bullet')
    p_f1_1.add_run('Auditoría y Corrección de Core Web Vitals: ').bold = True
    p_f1_1.add_run('Optimización de la carga asíncrona de videos en portada, retraso de scripts de terceros (Hotjar, GTM) para mejorar drásticamente el LCP y evitar penalizaciones de velocidad de Google.')
    p_f1_2 = doc.add_paragraph(style='List Bullet')
    p_f1_2.add_run('Reestructuración de Screaming Frog: ').bold = True
    p_f1_2.add_run('Escaneo total del sitio, arreglo de enlaces rotos (404), canónicas y corrección de la jerarquía H1-H3 para clarificar la arquitectura de los tours.')

    add_heading(doc, 'Fase 2: Answer Engine Optimization (AEO) - Semanas 3 y 4', 2)
    p_f2_1 = doc.add_paragraph(style='List Bullet')
    p_f2_1.add_run('Inyección de Schema Orgánico Avanzado: ').bold = True
    p_f2_1.add_run('Configuración en RankMath Pro de los esquemas FAQPage, HowTo y Product en los paquetes de viaje para ganar los Rich Snippets de Google SGE.')
    p_f2_2 = doc.add_paragraph(style='List Bullet')
    p_f2_2.add_run('Matriz "AlsoAsked": ').bold = True
    p_f2_2.add_run('Rediseño de las páginas de FAQs basándonos en consultas semánticas reales de viajeros de lujo.')

    add_heading(doc, 'Fase 3: Generative Engine Optimization (GEO) - Semanas 5 a 8', 2)
    p_f3_1 = doc.add_paragraph(style='List Bullet')
    p_f3_1.add_run('Desarrollo del Archivo llms.txt: ').bold = True
    p_f3_1.add_run('Creación e inyección técnica del archivo maestro para guiar a los bots de OpenAI y Anthropic sobre los valores diferenciadores de Auri Peru.')
    p_f3_2 = doc.add_paragraph(style='List Bullet')
    p_f3_2.add_run('Optimización de PR y Entidades: ').bold = True
    p_f3_2.add_run('Auditoría de menciones externas para aumentar la "autoridad" de la marca frente a los algoritmos de IA.')

    add_heading(doc, '3. Presupuesto y Cronograma de Inversión', 1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fase'
    hdr_cells[1].text = 'Entregable'
    hdr_cells[2].text = 'Tiempo'
    hdr_cells[3].text = 'Inversión (USD)'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Fase 1: SEO Técnico y CWV'
    row_cells[1].text = 'Mejora de PageSpeed (>80 Móvil), Reporte Screaming Frog 100% limpio.'
    row_cells[2].text = '2 Semanas'
    row_cells[3].text = '$ 1,500.00'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Fase 2: AEO & Schema'
    row_cells[1].text = 'Inyección de Schemas FAQPage y Hub de Respuestas.'
    row_cells[2].text = '2 Semanas'
    row_cells[3].text = '$ 1,200.00'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Fase 3: GEO y Optimización para IA'
    row_cells[1].text = 'Creación de llms.txt, Bing Webmaster y Pruebas ChatGPT.'
    row_cells[2].text = '4 Semanas'
    row_cells[3].text = '$ 2,800.00'

    row_cells = table.add_row().cells
    row_cells[0].paragraphs[0].add_run('TOTAL PROYECTO').bold = True
    row_cells[1].text = '-'
    row_cells[2].paragraphs[0].add_run('2 Meses').bold = True
    row_cells[3].paragraphs[0].add_run('$ 5,500.00').bold = True

    add_heading(doc, '4. Retorno de Inversión (ROI) Esperado', 1)
    doc.add_paragraph('Al completar esta consultoría, Auri Peru no solo mejorará su posicionamiento en Google tradicional, sino que se posicionará como una "Agencia Validada" dentro del ecosistema de Inteligencia Artificial (ChatGPT, Gemini), asegurando visibilidad ante viajeros de lujo globales, capturando leads cualificados de altísimo valor que hoy están absorbiendo Aracari o Kuoda.')

    doc.save(r'C:\Users\ASUS\OneDrive\VARIOS\Documentos\GPTS IA\BIOVET AI\Lifextreme-Web-AI\data\Propuesta_Comercial_AuriPeru.docx')

crear_auditoria()
crear_propuesta()
print("Archivos de Word generados exitosamente.")
